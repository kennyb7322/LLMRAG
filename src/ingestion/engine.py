"""
DATA INGESTION MODULE
=====================
Pulls data from multiple sources: filesystem, APIs, databases, S3, web crawlers.
Handles parsing, cleaning, deduplication, chunking, enrichment, and PII masking.
"""

import os
import re
import hashlib
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional, Generator
from dataclasses import dataclass, field
from datetime import datetime, timezone

from src.utils.logger import log


# ── Data Models ─────────────────────────────────────────────────────────────

@dataclass
class Document:
    """Represents an ingested document with metadata."""
    id: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    source: str = ""
    source_type: str = ""
    chunks: List["Chunk"] = field(default_factory=list)
    created_at: str = field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )

    def __post_init__(self):
        if not self.id:
            self.id = hashlib.sha256(self.content.encode()).hexdigest()[:16]


@dataclass
class Chunk:
    """A chunk of a document, ready for embedding."""
    id: str
    content: str
    document_id: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    index: int = 0


# ── Loaders ─────────────────────────────────────────────────────────────────

class FileSystemLoader:
    """Load documents from local filesystem."""

    SUPPORTED = {".txt", ".md", ".html", ".csv", ".json", ".pdf", ".docx"}

    def __init__(self, base_path: str, formats: List[str] = None, recursive: bool = True):
        self.base_path = Path(base_path)
        self.formats = [f".{f.lstrip('.')}" for f in (formats or [])] or list(self.SUPPORTED)
        self.recursive = recursive

    def load(self) -> Generator[Document, None, None]:
        """Yield documents from the filesystem."""
        pattern = "**/*" if self.recursive else "*"
        for filepath in self.base_path.glob(pattern):
            if filepath.is_file() and filepath.suffix.lower() in self.formats:
                try:
                    content = self._read_file(filepath)
                    if content and len(content.strip()) > 0:
                        yield Document(
                            id="",
                            content=content,
                            source=str(filepath),
                            source_type="filesystem",
                            metadata={
                                "filename": filepath.name,
                                "extension": filepath.suffix,
                                "size_bytes": filepath.stat().st_size,
                                "modified": datetime.fromtimestamp(
                                    filepath.stat().st_mtime, tz=timezone.utc
                                ).isoformat(),
                            },
                        )
                        log.info(f"Loaded: {filepath.name}")
                except Exception as e:
                    log.error(f"Failed to load {filepath}: {e}")

    def _read_file(self, filepath: Path) -> str:
        """Read file content based on type."""
        ext = filepath.suffix.lower()
        if ext in {".txt", ".md", ".html", ".csv", ".json"}:
            return filepath.read_text(encoding="utf-8", errors="replace")
        elif ext == ".pdf":
            return self._read_pdf(filepath)
        elif ext == ".docx":
            return self._read_docx(filepath)
        return ""

    @staticmethod
    def _read_pdf(filepath: Path) -> str:
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(filepath))
            text = "\n\n".join(page.get_text() for page in doc)
            doc.close()
            return text
        except ImportError:
            log.warning("PyMuPDF not installed — pip install pymupdf")
            return ""

    @staticmethod
    def _read_docx(filepath: Path) -> str:
        try:
            import docx
            doc = docx.Document(str(filepath))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            log.warning("python-docx not installed — pip install python-docx")
            return ""


class APILoader:
    """Load documents from REST APIs."""

    def __init__(self, endpoint: str, auth_type: str = "bearer",
                 auth_token: str = "", batch_size: int = 100):
        self.endpoint = endpoint
        self.auth_type = auth_type
        self.auth_token = auth_token or os.environ.get("API_AUTH_TOKEN", "")
        self.batch_size = batch_size

    def load(self) -> Generator[Document, None, None]:
        if not self.endpoint:
            log.info("No API endpoint configured — skipping API loader.")
            return
        try:
            import requests
            headers = self._build_headers()
            offset = 0
            while True:
                resp = requests.get(
                    self.endpoint,
                    headers=headers,
                    params={"limit": self.batch_size, "offset": offset},
                    timeout=30,
                )
                resp.raise_for_status()
                data = resp.json()
                items = data if isinstance(data, list) else data.get("results", [])
                if not items:
                    break
                for item in items:
                    content = item.get("content") or item.get("text") or str(item)
                    yield Document(
                        id="",
                        content=content,
                        source=self.endpoint,
                        source_type="api",
                        metadata={"raw_keys": list(item.keys()) if isinstance(item, dict) else []},
                    )
                offset += self.batch_size
                if len(items) < self.batch_size:
                    break
        except ImportError:
            log.warning("requests not installed — pip install requests")
        except Exception as e:
            log.error(f"API loader error: {e}")

    def _build_headers(self) -> dict:
        h = {"Content-Type": "application/json"}
        if self.auth_type == "bearer":
            h["Authorization"] = f"Bearer {self.auth_token}"
        elif self.auth_type == "api_key":
            h["X-API-Key"] = self.auth_token
        return h


class DatabaseLoader:
    """Load documents from SQL databases."""

    def __init__(self, connection_string: str, query: str, driver: str = "postgresql"):
        self.connection_string = connection_string
        self.query = query
        self.driver = driver

    def load(self) -> Generator[Document, None, None]:
        if not self.connection_string or not self.query:
            log.info("No database configured — skipping DB loader.")
            return
        try:
            import sqlalchemy
            engine = sqlalchemy.create_engine(self.connection_string)
            with engine.connect() as conn:
                result = conn.execute(sqlalchemy.text(self.query))
                for row in result:
                    row_dict = dict(row._mapping)
                    content = " ".join(str(v) for v in row_dict.values() if v)
                    yield Document(
                        id="",
                        content=content,
                        source=self.driver,
                        source_type="database",
                        metadata={"columns": list(row_dict.keys())},
                    )
        except ImportError:
            log.warning("sqlalchemy not installed — pip install sqlalchemy")
        except Exception as e:
            log.error(f"Database loader error: {e}")


class S3Loader:
    """Load documents from AWS S3."""

    def __init__(self, bucket: str, prefix: str = "", region: str = "us-east-1"):
        self.bucket = bucket
        self.prefix = prefix
        self.region = region

    def load(self) -> Generator[Document, None, None]:
        if not self.bucket:
            log.info("No S3 bucket configured — skipping S3 loader.")
            return
        try:
            import boto3
            s3 = boto3.client("s3", region_name=self.region)
            paginator = s3.get_paginator("list_objects_v2")
            for page in paginator.paginate(Bucket=self.bucket, Prefix=self.prefix):
                for obj in page.get("Contents", []):
                    key = obj["Key"]
                    resp = s3.get_object(Bucket=self.bucket, Key=key)
                    body = resp["Body"].read().decode("utf-8", errors="replace")
                    yield Document(
                        id="",
                        content=body,
                        source=f"s3://{self.bucket}/{key}",
                        source_type="s3",
                        metadata={"key": key, "size": obj["Size"]},
                    )
        except ImportError:
            log.warning("boto3 not installed — pip install boto3")
        except Exception as e:
            log.error(f"S3 loader error: {e}")


# ── Preprocessing ───────────────────────────────────────────────────────────

class Preprocessor:
    """Clean, deduplicate, mask PII."""

    def __init__(self, config: dict):
        self.cleaning = config.get("cleaning", {})
        self.dedup = config.get("deduplication", {})
        self.pii = config.get("pii_masking", {})
        self._seen_hashes = set()

    def process(self, doc: Document) -> Optional[Document]:
        """Clean, deduplicate, and mask a document. Returns None if duplicate."""
        # Clean
        text = doc.content
        if self.cleaning.get("remove_html"):
            text = re.sub(r"<[^>]+>", " ", text)
        if self.cleaning.get("normalize_whitespace"):
            text = re.sub(r"\s+", " ", text).strip()
        
        min_len = self.cleaning.get("min_content_length", 50)
        if len(text) < min_len:
            return None

        # Dedup
        if self.dedup.get("enabled"):
            content_hash = hashlib.md5(text.encode()).hexdigest()
            if content_hash in self._seen_hashes:
                log.info(f"Duplicate skipped: {doc.source}")
                return None
            self._seen_hashes.add(content_hash)

        # PII masking
        if self.pii.get("enabled"):
            text = self._mask_pii(text)

        doc.content = text
        return doc

    def _mask_pii(self, text: str) -> str:
        replacement = self.pii.get("replacement", "redacted")
        tag = f"[{replacement.upper()}]"
        # SSN
        text = re.sub(r"\b\d{3}-\d{2}-\d{4}\b", tag, text)
        # Credit card
        text = re.sub(r"\b\d{4}[\s-]?\d{4}[\s-]?\d{4}[\s-]?\d{4}\b", tag, text)
        # Email
        text = re.sub(r"\b[\w.+-]+@[\w-]+\.[\w.]+\b", tag, text)
        # Phone
        text = re.sub(r"\b(\+?1?[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b", tag, text)
        return text


# ── Chunking ────────────────────────────────────────────────────────────────

class Chunker:
    """Split documents into chunks for embedding."""

    def __init__(self, config: dict):
        self.strategy = config.get("strategy", "recursive")
        self.chunk_size = config.get("chunk_size", 512)
        self.overlap = config.get("chunk_overlap", 50)
        self.separators = config.get("separators", ["\n\n", "\n", ". ", " "])

    def chunk(self, doc: Document) -> List[Chunk]:
        """Split a document into chunks."""
        if self.strategy == "fixed":
            texts = self._fixed_chunks(doc.content)
        elif self.strategy == "sentence":
            texts = self._sentence_chunks(doc.content)
        elif self.strategy == "semantic":
            texts = self._semantic_chunks(doc.content)
        else:  # recursive
            texts = self._recursive_chunks(doc.content, self.separators)

        chunks = []
        for i, text in enumerate(texts):
            if text.strip():
                chunks.append(Chunk(
                    id=f"{doc.id}_chunk_{i:04d}",
                    content=text.strip(),
                    document_id=doc.id,
                    metadata={**doc.metadata, "chunk_index": i, "total_chunks": len(texts)},
                    index=i,
                ))
        return chunks

    def _fixed_chunks(self, text: str) -> List[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunks.append(text[start:end])
            start = end - self.overlap
        return chunks

    def _sentence_chunks(self, text: str) -> List[str]:
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks, current = [], ""
        for sent in sentences:
            if len(current) + len(sent) > self.chunk_size and current:
                chunks.append(current)
                # Overlap: keep last sentence
                overlap_sents = current.split(". ")
                current = ". ".join(overlap_sents[-1:]) + " " + sent if overlap_sents else sent
            else:
                current = (current + " " + sent).strip()
        if current.strip():
            chunks.append(current)
        return chunks

    def _semantic_chunks(self, text: str) -> List[str]:
        """Semantic chunking: split on double newlines first, then merge small paragraphs."""
        paragraphs = re.split(r'\n\n+', text)
        chunks, current = [], ""
        for para in paragraphs:
            if len(current) + len(para) > self.chunk_size and current:
                chunks.append(current)
                current = para
            else:
                current = (current + "\n\n" + para).strip()
        if current.strip():
            chunks.append(current)
        return chunks

    def _recursive_chunks(self, text: str, separators: List[str]) -> List[str]:
        """Recursively split text using a hierarchy of separators."""
        if not separators or len(text) <= self.chunk_size:
            return [text] if text.strip() else []
        
        sep = separators[0]
        parts = text.split(sep)
        chunks = []
        current = ""
        
        for part in parts:
            candidate = (current + sep + part).strip() if current else part.strip()
            if len(candidate) > self.chunk_size and current:
                chunks.append(current)
                current = part.strip()
            else:
                current = candidate
        if current.strip():
            chunks.append(current)

        # Recursively split any oversized chunks
        final = []
        for c in chunks:
            if len(c) > self.chunk_size and len(separators) > 1:
                final.extend(self._recursive_chunks(c, separators[1:]))
            else:
                final.append(c)
        return final


# ── Enrichment ──────────────────────────────────────────────────────────────

class Enricher:
    """Add metadata, entities, keywords, summaries to chunks."""

    def __init__(self, config: dict):
        self.extract_metadata = config.get("extract_metadata", True)
        self.extract_entities = config.get("extract_entities", False)
        self.generate_keywords = config.get("generate_keywords", True)
        self.generate_summaries = config.get("generate_summaries", False)
        self.add_timestamps = config.get("add_timestamps", True)

    def enrich(self, chunk: Chunk) -> Chunk:
        """Enrich a chunk with additional metadata."""
        if self.add_timestamps:
            chunk.metadata["enriched_at"] = datetime.now(timezone.utc).isoformat()

        if self.generate_keywords:
            chunk.metadata["keywords"] = self._extract_keywords(chunk.content)

        if self.extract_entities:
            chunk.metadata["entities"] = self._extract_entities(chunk.content)

        chunk.metadata["word_count"] = len(chunk.content.split())
        chunk.metadata["char_count"] = len(chunk.content)

        return chunk

    @staticmethod
    def _extract_keywords(text: str, top_n: int = 10) -> List[str]:
        """Simple TF-based keyword extraction."""
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
        # Remove common stopwords
        stops = {"this", "that", "with", "from", "have", "been", "will", "would",
                 "could", "should", "their", "there", "about", "which", "when",
                 "what", "where", "they", "them", "then", "than", "into", "also",
                 "some", "more", "other", "were", "your", "each", "very", "most"}
        words = [w for w in words if w not in stops]
        freq = {}
        for w in words:
            freq[w] = freq.get(w, 0) + 1
        return sorted(freq, key=freq.get, reverse=True)[:top_n]

    @staticmethod
    def _extract_entities(text: str) -> List[Dict[str, str]]:
        """Basic named entity extraction using regex patterns."""
        entities = []
        # Capitalized multi-word names
        for match in re.finditer(r'\b([A-Z][a-z]+ (?:[A-Z][a-z]+ ?)+)', text):
            entities.append({"text": match.group().strip(), "type": "NAME"})
        # Dates
        for match in re.finditer(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', text):
            entities.append({"text": match.group(), "type": "DATE"})
        return entities[:20]  # Limit


# ── Main Ingestion Engine ───────────────────────────────────────────────────

class IngestionEngine:
    """Orchestrates the full ingestion pipeline."""

    def __init__(self, config: dict):
        self.config = config
        self.preprocessor = Preprocessor(config.get("preprocessing", {}))
        self.chunker = Chunker(config.get("chunking", {}))
        self.enricher = Enricher(config.get("enrichment", {}))

    def run(self) -> List[Chunk]:
        """Execute the full ingestion pipeline. Returns enriched chunks."""
        log.info("═══ INGESTION PIPELINE STARTED ═══")
        
        documents = list(self._load_all_sources())
        log.info(f"Loaded {len(documents)} raw documents")

        # Preprocess
        cleaned = []
        for doc in documents:
            result = self.preprocessor.process(doc)
            if result:
                cleaned.append(result)
        log.info(f"After preprocessing: {len(cleaned)} documents")

        # Chunk
        all_chunks = []
        for doc in cleaned:
            chunks = self.chunker.chunk(doc)
            doc.chunks = chunks
            all_chunks.extend(chunks)
        log.info(f"Generated {len(all_chunks)} chunks")

        # Enrich
        enriched = [self.enricher.enrich(c) for c in all_chunks]
        log.info(f"Enriched {len(enriched)} chunks")
        log.info("═══ INGESTION PIPELINE COMPLETE ═══")

        return enriched

    def _load_all_sources(self) -> Generator[Document, None, None]:
        """Load from all configured sources."""
        for source_cfg in self.config.get("sources", []):
            src_type = source_cfg.get("type", "")
            
            if src_type == "filesystem":
                loader = FileSystemLoader(
                    base_path=source_cfg.get("path", "./data/documents"),
                    formats=source_cfg.get("formats"),
                    recursive=source_cfg.get("recursive", True),
                )
                yield from loader.load()
            
            elif src_type == "api":
                loader = APILoader(
                    endpoint=source_cfg.get("endpoint", ""),
                    auth_type=source_cfg.get("auth_type", "bearer"),
                    batch_size=source_cfg.get("batch_size", 100),
                )
                yield from loader.load()
            
            elif src_type == "database":
                loader = DatabaseLoader(
                    connection_string=source_cfg.get("connection_string", ""),
                    query=source_cfg.get("query", ""),
                    driver=source_cfg.get("driver", "postgresql"),
                )
                yield from loader.load()
            
            elif src_type == "s3":
                loader = S3Loader(
                    bucket=source_cfg.get("bucket", ""),
                    prefix=source_cfg.get("prefix", ""),
                    region=source_cfg.get("region", "us-east-1"),
                )
                yield from loader.load()
            
            elif src_type == "web_crawler":
                log.info("Web crawler source detected — use tools/web_crawler.py")
